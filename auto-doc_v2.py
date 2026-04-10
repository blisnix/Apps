import sys
import os
import platform
import subprocess
import threading
import importlib
import datetime
import tkinter as tk
from tkinter import ttk, messagebox
from functools import lru_cache

OS         = platform.system()
PY_VERSION = sys.version_info
OS_NAMES   = {"Windows": "Windows", "Darwin": "macOS", "Linux": "Linux"}

_UI_FONT = {"Windows": "Segoe UI", "Darwin": "Helvetica Neue"}.get(OS, "Ubuntu")

THEME = {
    "BG":     "#0d1117",
    "BG2":    "#161b22",
    "BG3":    "#1c2333",
    "ACCENT": "#58a6ff",
    "GREEN":  "#3fb950",
    "RED":    "#f85149",
    "YELLOW": "#d29922",
    "GREY":   "#8b949e",
    "FG":     "#e6edf3",
    "BORDER": "#30363d",
    "FONT":   _UI_FONT,
}

VENV_DIR    = os.path.join(os.path.expanduser("~"), ".autodoc_venv")
VENV_PYTHON = os.path.join(VENV_DIR, "Scripts" if OS == "Windows" else "bin", "python")

def _in_venv():
    return os.path.abspath(sys.executable).startswith(os.path.abspath(VENV_DIR))

def _ensure_venv_and_relaunch():
    if _in_venv():
        return
    if not os.path.isfile(VENV_PYTHON):
        print("Creating virtual environment at", VENV_DIR)
        subprocess.check_call([sys.executable, "-m", "venv", VENV_DIR])
    print("Relaunching inside virtual environment...")
    if OS == "Windows":
        subprocess.Popen([VENV_PYTHON] + sys.argv); sys.exit(0)
    else:
        os.execv(VENV_PYTHON, [VENV_PYTHON] + sys.argv); sys.exit(0)

_ensure_venv_and_relaunch()

def _can_import(mod):
    try: importlib.import_module(mod); return True
    except ImportError: return False

@lru_cache(maxsize=16)
def _cmd_exists(cmd):
    try:
        return subprocess.run(["where" if OS == "Windows" else "which", cmd],
                              capture_output=True, timeout=3).returncode == 0
    except Exception: return False

ALL_PREREQS = [
    {"name":"Python 3.8+","pip":None,"purpose":"Core runtime",
     "platform":["all"],"check":lambda:PY_VERSION>=(3,8),
     "fix":"Download Python 3.8+ from https://python.org/downloads"},
    {"name":"Tkinter (GUI)","pip":None,"purpose":"Provides the graphical user interface",
     "platform":["all"],"check":lambda:_can_import("tkinter"),
     "fix":{"Linux":"sudo apt install python3-tk","Darwin":"Reinstall Python from python.org",
            "Windows":"Reinstall Python, check 'tcl/tk' option"}},
    {"name":"Pillow (PIL)","pip":"pillow","purpose":"Takes and processes screenshots",
     "platform":["all"],"check":lambda:_can_import("PIL"),"fix":"pip install pillow"},
    {"name":"python-docx","pip":"python-docx","purpose":"Creates Word (.docx) documents",
     "platform":["all"],"check":lambda:_can_import("docx"),"fix":"pip install python-docx"},
    {"name":"pynput","pip":"pynput","purpose":"Reads mouse position for multi-monitor support",
     "platform":["all"],"check":lambda:_can_import("pynput"),"fix":"pip install pynput"},
    {"name":"screeninfo","pip":"screeninfo","purpose":"Detects monitor layout",
     "platform":["all"],"check":lambda:_can_import("screeninfo"),"fix":"pip install screeninfo"},
    {"name":"pygetwindow","pip":"pygetwindow","purpose":"Detects active window (Windows)",
     "platform":["Windows"],"check":lambda:_can_import("pygetwindow"),"fix":"pip install pygetwindow"},
    {"name":"pyobjc-Cocoa","pip":"pyobjc-framework-Cocoa","purpose":"Detects active app (macOS)",
     "platform":["Darwin"],"check":lambda:_can_import("AppKit"),"fix":"pip install pyobjc-framework-Cocoa"},
    {"name":"xdotool (system)","pip":None,"purpose":"Detects active window (Linux)",
     "platform":["Linux"],"check":lambda:_cmd_exists("xdotool"),
     "fix":"sudo apt install xdotool"},
]

def get_platform_prereqs():
    return [p for p in ALL_PREREQS if "all" in p["platform"] or OS in p["platform"]]

def get_fix_msg(prereq):
    fix = prereq.get("fix", "")
    return fix.get(OS, next(iter(fix.values()))) if isinstance(fix, dict) else str(fix)

# === SHARED UI HELPERS ===

class ThemedButton(tk.Frame):
    """Label-based button — works on macOS where tk.Button ignores bg/fg."""
    def __init__(self, parent, text, cmd, bg, hover_bg, width=None):
        super().__init__(parent, bg=bg)
        self._cmd, self._bg, self._hover_bg = cmd, bg, hover_bg
        kw = dict(text=text, font=(THEME["FONT"], 10, "bold"),
                  bg=bg, fg=THEME["FG"], padx=16, pady=9)
        if width:
            kw["width"] = width
        self._lbl = tk.Label(self, **kw)
        self._lbl.pack()
        for w in (self, self._lbl):
            w.bind("<Button-1>", self._click)
            w.bind("<Enter>",    self._enter)
            w.bind("<Leave>",    self._leave)

    def _click(self, _): self._cmd()
    def _enter(self, _): self._set_bg(self._hover_bg)
    def _leave(self, _): self._set_bg(self._bg)

    def _set_bg(self, color):
        tk.Frame.config(self, bg=color)
        self._lbl.config(bg=color)

    def config(self, **kw):
        state = kw.pop("state", None)
        text  = kw.pop("text", None)
        if state == "disabled":
            self._lbl.config(fg=THEME["GREY"])
            for w in (self, self._lbl):
                w.unbind("<Button-1>")
        elif state == "normal":
            self._lbl.config(fg=THEME["FG"])
            for w in (self, self._lbl):
                w.bind("<Button-1>", self._click)
        if text is not None:
            self._lbl.config(text=text)
        if kw:
            tk.Frame.config(self, **kw)

def make_button(parent, text, cmd, bg, hover_bg, width=None):
    return ThemedButton(parent, text, cmd, bg, hover_bg, width)

def make_section_label(parent, text):
    f = tk.Frame(parent, bg=THEME["BG"], padx=14); f.pack(fill="x", pady=(10, 3))
    tk.Label(f, text=f"  {text}  ", font=(THEME["FONT"], 8, "bold"),
             fg=THEME["ACCENT"], bg=THEME["BG"]).pack(side="left")
    tk.Frame(f, bg=THEME["BORDER"], height=1).pack(side="left", fill="x", expand=True, pady=5)

# === LAUNCHER ===

class Launcher:
    def __init__(self, root):
        self.root = root
        root.title("Auto-Doc -- Setup & Launcher")
        root.geometry("780x660")
        root.configure(bg=THEME["BG"])
        root.resizable(True, True)
        self.prereqs = get_platform_prereqs()
        self.icon_labels = {}
        self.status_labels = {}
        self._build_ui()
        root.after(300, self._run_checks)

    def _build_ui(self):
        T = THEME
        hdr = tk.Frame(self.root, bg=T["BG2"], pady=22); hdr.pack(fill="x")
        tk.Label(hdr, text="AUTO-DOC", font=(T["FONT"], 22, "bold"),
                 fg=T["ACCENT"], bg=T["BG2"]).pack()
        tk.Label(hdr, text="Process Documentation Tool", font=(T["FONT"], 10),
                 fg=T["GREY"], bg=T["BG2"]).pack(pady=(2, 8))
        br = tk.Frame(hdr, bg=T["BG2"]); br.pack()
        oc = {"Windows":"#1d4ed8","Darwin":"#7c3aed","Linux":"#15803d"}
        tk.Label(br, text=f"  {OS_NAMES.get(OS,OS)}  ", font=(T["FONT"],8,"bold"),
                 fg="white", bg=oc.get(OS,"#374151"), padx=8, pady=3).pack(side="left")
        tk.Label(br, text=f"  Python {PY_VERSION.major}.{PY_VERSION.minor}.{PY_VERSION.micro}  ·  Setup & Launcher",
                 font=(T["FONT"],9), fg=T["GREY"], bg=T["BG2"]).pack(side="left", padx=8)
        intro = tk.Frame(self.root, bg=T["BG"], padx=20, pady=12); intro.pack(fill="x")
        tk.Label(intro, text="Checks all prerequisites needed to run Auto-Doc.\nMissing packages can be installed automatically.",
                 font=(T["FONT"],9), fg=T["GREY"], bg=T["BG"], justify="left").pack(anchor="w")
        make_section_label(self.root, "PREREQUISITE STATUS")
        lo = tk.Frame(self.root, bg=T["BG"], padx=14); lo.pack(fill="both", expand=True)
        canvas = tk.Canvas(lo, bg=T["BG3"], highlightthickness=0, bd=0)
        sb = tk.Scrollbar(lo, orient="vertical", command=canvas.yview)
        self.pf = tk.Frame(canvas, bg=T["BG3"])
        self.pf.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
        canvas.create_window((0,0), window=self.pf, anchor="nw")
        canvas.configure(yscrollcommand=sb.set)
        canvas.pack(side="left", fill="both", expand=True); sb.pack(side="right", fill="y")
        for i, p in enumerate(self.prereqs): self._build_row(self.pf, p, i)
        make_section_label(self.root, "INSTALL LOG")
        lf = tk.Frame(self.root, bg=T["BG"], padx=14); lf.pack(fill="x")
        ls = tk.Scrollbar(lf); ls.pack(side="right", fill="y")
        self.log_box = tk.Text(lf, height=5, font=(T["FONT"],9), bg=T["BG2"], fg=T["ACCENT"],
                               insertbackground=T["ACCENT"], relief="flat", bd=0,
                               yscrollcommand=ls.set, state="disabled", wrap="word")
        self.log_box.pack(fill="x"); ls.config(command=self.log_box.yview)
        for tag, color in [("ok",T["GREEN"]),("err",T["RED"]),("info",T["ACCENT"]),("warn",T["YELLOW"])]:
            self.log_box.tag_config(tag, foreground=color)
        bb = tk.Frame(self.root, bg=T["BG2"], pady=12); bb.pack(fill="x", side="bottom")
        self.install_btn = make_button(bb, "Install All Missing", self._install_all, "#1a3a1a", "#2e7d32", width=24)
        self.install_btn.pack(side="left", padx=(16,8))
        make_button(bb, "Re-check", self._run_checks, "#1b2838", "#263545", width=16).pack(side="left", padx=(0,8))
        self.launch_btn = make_button(bb, "Launch Application", self._launch_app, "#003d1f", "#00695c", width=24)
        self.launch_btn.pack(side="right", padx=(0,16)); self.launch_btn.config(state="disabled")
        self.overall_var = tk.StringVar(value="Checking prerequisites...")
        tk.Label(bb, textvariable=self.overall_var, font=(T["FONT"],9,"bold"),
                 fg=T["YELLOW"], bg=T["BG2"]).pack(side="left", padx=8)

    def _build_row(self, parent, prereq, idx):
        T = THEME; bg = T["BG3"] if idx % 2 == 0 else "#111c2b"
        row = tk.Frame(parent, bg=bg, padx=12, pady=8); row.pack(fill="x", pady=1)
        il = tk.Label(row, text="·", font=(T["FONT"],13), fg=T["YELLOW"], bg=bg, width=2)
        il.pack(side="left"); self.icon_labels[prereq["name"]] = il
        col = tk.Frame(row, bg=bg); col.pack(side="left", fill="x", expand=True)
        tk.Label(col, text=prereq["name"], font=(T["FONT"],10,"bold"), fg=T["FG"], bg=bg, anchor="w").pack(fill="x")
        tk.Label(col, text=prereq["purpose"], font=(T["FONT"],8), fg=T["GREY"], bg=bg, anchor="w").pack(fill="x")
        sl = tk.Label(row, text="Checking...", font=(T["FONT"],9), fg=T["YELLOW"], bg=bg, width=20, anchor="e")
        sl.pack(side="right"); self.status_labels[prereq["name"]] = sl

    def _log(self, msg, tag="info"):
        ts = datetime.datetime.now().strftime("%H:%M:%S")
        self.log_box.config(state="normal")
        self.log_box.insert("end", f"[{ts}] {msg}\n", tag)
        self.log_box.see("end")
        self.log_box.config(state="disabled")

    def _log_later(self, msg, tag="info"):
        self.root.after(0, lambda: self._log(msg, tag))

    def _set_row(self, name, icon, status, ifg, sfg):
        self.icon_labels[name].config(text=icon, fg=ifg)
        self.status_labels[name].config(text=status, fg=sfg)

    def _run_checks(self):
        self.overall_var.set("Checking..."); self.launch_btn.config(state="disabled")
        threading.Thread(target=self._check_thread, daemon=True).start()

    def _check_thread(self):
        T = THEME; all_ok = True; missing = []
        for p in self.prereqs:
            n = p["name"]
            self.root.after(0, lambda n=n: (self.icon_labels[n].config(text="..."),
                                            self.status_labels[n].config(text="Checking...")))
            try: ok = p["check"]()
            except: ok = False
            if ok:
                self.root.after(0, lambda n=n: self._set_row(n,"✓","Installed",T["GREEN"],T["GREEN"]))
                self._log_later(f"✓  {n}", "ok")
            else:
                all_ok = False
                if p.get("pip"):
                    missing.append(p)
                    self.root.after(0, lambda n=n: self._set_row(n,"✗","Missing",T["RED"],T["YELLOW"]))
                    self._log_later(f"✗  {n} — can auto-install", "warn")
                else:
                    self.root.after(0, lambda n=n: self._set_row(n,"⚠","Manual fix",T["YELLOW"],T["YELLOW"]))
                    self._log_later(f"⚠  {n} — {get_fix_msg(p)}", "warn")
        if all_ok: self.root.after(0, self._on_all_ok)
        elif missing: self.root.after(0, lambda: self.overall_var.set(f"{len(missing)} package(s) missing"))
        else: self.root.after(0, lambda: self.overall_var.set("Manual fixes required"))

    def _on_all_ok(self):
        self.overall_var.set("All prerequisites satisfied — ready to launch!")
        self.launch_btn.config(state="normal", bg="#003d1f")
        self._log("All checks passed.", "ok")

    def _install_all(self):
        self.install_btn.config(state="disabled", text="Installing...")
        self._log("Starting installation...", "info")
        threading.Thread(target=self._install_thread, daemon=True).start()

    def _mark_result(self, p, ok):
        T = THEME
        if ok:
            self._log_later(f"✓  {p['name']}", "ok")
            self.root.after(0, lambda n=p["name"]: self._set_row(n, "✓", "Installed", T["GREEN"], T["GREEN"]))
        else:
            self._log_later(f"✗  {p['name']} failed", "err")
            self.root.after(0, lambda n=p["name"]: self._set_row(n, "✗", "Failed", T["RED"], T["RED"]))

    def _install_thread(self):
        T = THEME
        missing = [p for p in self.prereqs if p.get("pip") and not p["check"]()]
        if not missing:
            self._log_later("Nothing to install.", "ok")
            self.root.after(0, lambda: self.install_btn.config(state="normal", text="Install All Missing"))
            return
        pkgs = [p["pip"] for p in missing]
        for p in missing:
            self.root.after(0, lambda n=p["name"]: self._set_row(n, "⟳", "Installing...", T["YELLOW"], T["YELLOW"]))
        self._log_later(f"Batch installing: {', '.join(pkgs)}", "info")
        try:
            r = subprocess.run([sys.executable, "-m", "pip", "install"] + pkgs + ["--quiet", "--no-warn-script-location"],
                               capture_output=True, text=True, timeout=300)
            if r.returncode == 0:
                for p in missing:
                    self._mark_result(p, ok=True)
            else:
                self._log_later("Batch failed, trying individually...", "warn")
                for p in missing:
                    try:
                        r2 = subprocess.run([sys.executable, "-m", "pip", "install", p["pip"], "--quiet"],
                                            capture_output=True, text=True, timeout=180)
                        self._mark_result(p, ok=r2.returncode == 0)
                    except Exception as e:
                        self._log_later(f"X  {p['name']} error: {e}", "err")
        except Exception as e:
            self._log_later(f"Install error: {e}", "err")
        self.root.after(0, lambda: self.install_btn.config(state="normal", text="Install All Missing"))
        self.root.after(600, self._run_checks)

    def _launch_app(self):
        self._log("Launching...", "info"); self.root.after(200, self._do_launch)

    def _do_launch(self):
        self.root.destroy()
        if OS == "Windows": subprocess.Popen([sys.executable, __file__, "--run-app"]); sys.exit(0)
        else: os.execv(sys.executable, [sys.executable, __file__, "--run-app"])

# === MAIN APPLICATION ===

def _run_main_app():
    import datetime, time, json, io, tempfile, atexit, shutil, math

    from PIL import ImageGrab, Image, ImageDraw, ImageFont, ImageTk
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from pynput import keyboard as pynput_keyboard
    from tkinter import messagebox, simpledialog, filedialog, colorchooser

    T = THEME

    # -- Cached mouse controller --
    _mc = [None]
    def get_mouse_pos():
        try:
            if _mc[0] is None:
                from pynput.mouse import Controller; _mc[0] = Controller()
            pos = _mc[0].position; return int(pos[0]), int(pos[1])
        except: return 0, 0

    # -- Cached monitors --
    _mon_cache = {"m": None, "t": 0}
    def _get_monitors():
        now = time.monotonic()
        if _mon_cache["m"] is None or now - _mon_cache["t"] > 5.0:
            try:
                from screeninfo import get_monitors; _mon_cache["m"] = get_monitors()
            except: _mon_cache["m"] = []
            _mon_cache["t"] = now
        return _mon_cache["m"]

    def get_monitor_at(mx, my):
        mons = _get_monitors()
        for m in mons:
            if m.x <= mx < m.x+m.width and m.y <= my < m.y+m.height:
                return m.x, m.y, m.width, m.height
        if mons: m = mons[0]; return m.x, m.y, m.width, m.height
        return None

    def annotate_cursor(img, cx, cy, mx=0, my=0):
        try:
            draw = ImageDraw.Draw(img)
            px, py = cx - mx, cy - my
            base = min(img.width, img.height); r = max(18, int(base*0.025)); th = max(2, int(base*0.003))
            R, W = (220,30,30,230), (255,255,255,180); hw = th+1
            draw.ellipse([px-r-hw,py-r-hw,px+r+hw,py+r+hw], outline=W, width=hw)
            draw.ellipse([px-r,py-r,px+r,py+r], outline=R, width=th)
            gap = r+4; arm = int(r*0.75)
            for dx, dy in [(1,0),(-1,0),(0,1),(0,-1)]:
                x1, y1 = px+dx*gap, py+dy*gap
                x2, y2 = px+dx*(gap+arm), py+dy*(gap+arm)
                draw.line([x1,y1,x2,y2], fill=W, width=th+2)
                draw.line([x1,y1,x2,y2], fill=R, width=th)
        except: pass
        return img

    def take_screenshot_at_mouse(highlight=True):
        mx, my = get_mouse_pos(); monitor = get_monitor_at(mx, my)
        img = None; mon_x = mon_y = 0
        try:
            if OS == "Darwin":
                tmp = tempfile.mktemp(suffix=".png")
                if monitor:
                    x,y,w,h = monitor; mon_x,mon_y = x,y
                    subprocess.run(["screencapture","-x","-R",f"{x},{y},{w},{h}",tmp], timeout=4)
                else: subprocess.run(["screencapture","-x",tmp], timeout=4)
                img = Image.open(tmp); os.unlink(tmp)
            else:
                kw = {}
                if monitor:
                    x,y,w,h = monitor; mon_x,mon_y = x,y; kw["bbox"] = (x,y,x+w,y+h)
                if OS == "Windows": kw["all_screens"] = True
                img = ImageGrab.grab(**kw)
        except:
            try: img = ImageGrab.grab()
            except: return None
        if img and highlight:
            img = img.convert("RGBA"); img = annotate_cursor(img, mx, my, mon_x, mon_y); img = img.convert("RGB")
        return img

    def open_file(path):
        try:
            if OS == "Windows": os.startfile(path)
            elif OS == "Darwin": subprocess.Popen(["open", path])
            else: subprocess.Popen(["xdg-open", path])
        except: pass

    # -- F-key listener --
    _app_ref = [None]
    def on_fkey(key):
        app = _app_ref[0]
        if app is None: return
        try:
            if key == pynput_keyboard.Key.f9: app.root.after(0, app._add_step_with_shot)
            elif key == pynput_keyboard.Key.f10: app.root.after(0, app._toggle_pause_capture)
            elif key == pynput_keyboard.Key.f11: app.root.after(0, app._export_word)
        except: pass
    try:
        fkey_listener = pynput_keyboard.Listener(on_press=on_fkey); fkey_listener.daemon = True; fkey_listener.start()
    except: fkey_listener = None

    # ===== SCREENSHOT EDITOR =====

    class ScreenshotEditor:
        """Full-window screenshot annotation editor with arrow, rect, freehand, text, highlight, crop."""
        TOOLS = ["arrow","rectangle","freehand","text","highlight","crop"]
        LABELS = {"arrow":"Arrow","rectangle":"Rectangle","freehand":"Freehand",
                  "text":"Text","highlight":"Highlight","crop":"Crop"}

        def __init__(self, parent, pil_image, on_save=None):
            self.on_save = on_save
            self.original = pil_image.copy()
            self.img = pil_image.copy().convert("RGBA")
            self.history = [self.img.copy()]
            self.tool = "arrow"
            self.color = "#ff0000"
            self.width = 3
            self._start = None
            self._freehand = []
            self._preview_id = None

            self.win = tk.Toplevel(parent)
            self.win.title("Screenshot Editor")
            self.win.geometry("1100x780")
            self.win.configure(bg=T["BG"])
            self.win.grab_set()
            self._build_ui()
            self._render()

        def _build_ui(self):
            tb = tk.Frame(self.win, bg=T["BG2"], pady=6); tb.pack(fill="x")
            tk.Label(tb, text=" Tool:", font=(T["FONT"],9), fg=T["GREY"], bg=T["BG2"]).pack(side="left", padx=(8,4))
            self.tool_var = tk.StringVar(value="arrow")
            for t in self.TOOLS:
                rb = tk.Radiobutton(tb, text=self.LABELS[t], variable=self.tool_var, value=t,
                    command=self._tool_changed, font=(T["FONT"],9,"bold"), fg=T["FG"], bg=T["BG2"],
                    selectcolor=T["BG3"], activebackground=T["BG3"], activeforeground=T["ACCENT"],
                    indicatoron=0, padx=10, pady=5, relief="flat", bd=0)
                rb.pack(side="left", padx=2)
            tk.Frame(tb, bg=T["GREY"], width=1).pack(side="left", fill="y", padx=8, pady=4)
            self.swatch = tk.Label(tb, text="  ", bg=self.color, width=3, relief="raised", bd=1)
            self.swatch.pack(side="left", padx=(4,2)); self.swatch.bind("<Button-1>", self._pick_color)
            tk.Label(tb, text="Color", font=(T["FONT"],8), fg=T["GREY"], bg=T["BG2"]).pack(side="left", padx=(0,8))
            tk.Label(tb, text="Width:", font=(T["FONT"],8), fg=T["GREY"], bg=T["BG2"]).pack(side="left")
            self.w_var = tk.IntVar(value=3)
            tk.Scale(tb, from_=1, to=12, orient="horizontal", variable=self.w_var, length=100,
                     showvalue=True, font=(T["FONT"],7), fg=T["FG"], bg=T["BG2"],
                     troughcolor=T["BG3"], highlightthickness=0, bd=0).pack(side="left", padx=(2,8))
            tk.Frame(tb, bg=T["GREY"], width=1).pack(side="left", fill="y", padx=4, pady=4)
            make_button(tb, "Undo", self._undo, "#1b2838", "#263545").pack(side="left", padx=4)
            make_button(tb, "Reset", self._reset, "#3b1515", "#5c1f1f").pack(side="left", padx=4)
            make_button(tb, "Save to Step", self._save_to_step, "#1a3a1a", "#2e7d32").pack(side="right", padx=(4,8))
            make_button(tb, "Export Image", self._export_image, "#1b2838", "#263545").pack(side="right", padx=4)

            cf = tk.Frame(self.win, bg=T["BG"]); cf.pack(fill="both", expand=True, padx=8, pady=8)
            self.hs = tk.Scrollbar(cf, orient="horizontal"); self.vs = tk.Scrollbar(cf, orient="vertical")
            self.canvas = tk.Canvas(cf, bg="#111111", highlightthickness=0,
                                    xscrollcommand=self.hs.set, yscrollcommand=self.vs.set, cursor="crosshair")
            self.hs.config(command=self.canvas.xview); self.vs.config(command=self.canvas.yview)
            self.vs.pack(side="right", fill="y"); self.hs.pack(side="bottom", fill="x")
            self.canvas.pack(side="left", fill="both", expand=True)
            self.canvas.bind("<ButtonPress-1>", self._press)
            self.canvas.bind("<B1-Motion>", self._drag)
            self.canvas.bind("<ButtonRelease-1>", self._release)
            self.status = tk.StringVar(value="Select a tool and draw on the screenshot")
            tk.Label(self.win, textvariable=self.status, font=(T["FONT"],9), fg=T["ACCENT"],
                     bg=T["BG2"], anchor="w", padx=12, pady=5).pack(fill="x", side="bottom")

        def _tool_changed(self):
            self.tool = self.tool_var.get()
            cursors = {"freehand":"pencil","text":"xterm"}
            self.canvas.config(cursor=cursors.get(self.tool, "crosshair"))

        def _pick_color(self, e=None):
            r = colorchooser.askcolor(color=self.color, title="Pick Color", parent=self.win)
            if r and r[1]: self.color = r[1]; self.swatch.config(bg=self.color)

        def _render(self):
            disp = self.img.copy().convert("RGB")
            self._tk_img = ImageTk.PhotoImage(disp)
            self.canvas.delete("all")
            self.canvas.create_image(0, 0, anchor="nw", image=self._tk_img)
            self.canvas.config(scrollregion=(0, 0, disp.width, disp.height))

        def _cc(self, e):
            return int(self.canvas.canvasx(e.x)), int(self.canvas.canvasy(e.y))

        def _push(self):
            self.history.append(self.img.copy())
            if len(self.history) > 30: self.history.pop(0)

        def _undo(self):
            if len(self.history) > 1:
                self.history.pop()
                self.img = self.history[-1].copy()
                self._render()
                self.status.set("Undo")
            else:
                self.status.set("Nothing to undo")

        def _reset(self):
            self.img = self.original.copy().convert("RGBA")
            self.history = [self.img.copy()]
            self._render()
            self.status.set("Reset to original")

        def _press(self, e):
            x, y = self._cc(e)
            self._start = (x, y)
            self._freehand = [(x, y)]
            self.width = self.w_var.get()

        def _drag(self, e):
            x, y = self._cc(e)
            if self.tool == "freehand":
                if self._freehand:
                    lx, ly = self._freehand[-1]
                    self.canvas.create_line(lx,ly,x,y, fill=self.color, width=self.width, tags="preview")
                self._freehand.append((x, y))
            elif self.tool in ("arrow","rectangle","highlight","crop"):
                if self._preview_id: self.canvas.delete(self._preview_id)
                sx, sy = self._start
                if self.tool == "arrow":
                    self._preview_id = self.canvas.create_line(sx,sy,x,y, fill=self.color, width=self.width, arrow=tk.LAST, tags="preview")
                elif self.tool == "rectangle":
                    self._preview_id = self.canvas.create_rectangle(sx,sy,x,y, outline=self.color, width=self.width, tags="preview")
                elif self.tool == "crop":
                    self._preview_id = self.canvas.create_rectangle(sx,sy,x,y, outline="#00e5ff", width=2, dash=(4,4), tags="preview")
                elif self.tool == "highlight":
                    self._preview_id = self.canvas.create_rectangle(sx,sy,x,y, outline="#ffeb3b", fill="#ffeb3b", stipple="gray25", tags="preview")

        def _release(self, e):
            x, y = self._cc(e); self.canvas.delete("preview")
            if self._start is None: return
            sx, sy = self._start; self._push()
            draw = ImageDraw.Draw(self.img); c = self.color; w = self.width

            if self.tool == "arrow":
                draw.line([sx,sy,x,y], fill=c, width=w)
                angle = math.atan2(y-sy, x-sx); hl = max(12, w*5); ha = math.pi/6
                for s in (1, -1):
                    hx = x - hl*math.cos(angle+s*ha); hy = y - hl*math.sin(angle+s*ha)
                    draw.line([x,y,int(hx),int(hy)], fill=c, width=w)
            elif self.tool == "rectangle":
                draw.rectangle([sx,sy,x,y], outline=c, width=w)
            elif self.tool == "freehand":
                if len(self._freehand) > 1: draw.line(self._freehand, fill=c, width=w, joint="curve")
            elif self.tool == "text":
                txt = simpledialog.askstring("Add Text", "Enter annotation:", parent=self.win)
                if txt:
                    fs = max(14, w*6)
                    try: font = ImageFont.truetype("arial.ttf", fs)
                    except:
                        try: font = ImageFont.truetype("/System/Library/Fonts/Helvetica.ttc", fs)
                        except: font = ImageFont.load_default()
                    draw.text((sx+2,sy+2), txt, fill="black", font=font)
                    draw.text((sx,sy), txt, fill=c, font=font)
                else: self.history.pop()
            elif self.tool == "highlight":
                ov = Image.new("RGBA", self.img.size, (0,0,0,0))
                ImageDraw.Draw(ov).rectangle([min(sx,x),min(sy,y),max(sx,x),max(sy,y)], fill=(255,255,0,80))
                self.img = Image.alpha_composite(self.img, ov); self.history[-1] = self.img.copy()
            elif self.tool == "crop":
                cx1,cy1,cx2,cy2 = min(sx,x),min(sy,y),max(sx,x),max(sy,y)
                if cx2-cx1 > 10 and cy2-cy1 > 10:
                    self.img = self.img.crop((cx1,cy1,cx2,cy2)); self.history[-1] = self.img.copy()
                    self.status.set(f"Cropped to {cx2-cx1}x{cy2-cy1}")

            self._start = None
            self._freehand = []
            self._preview_id = None
            self._render()

        def _save_to_step(self):
            if self.on_save: self.on_save(self.img.convert("RGB"))
            self.status.set("Saved to step"); self.win.after(600, self.win.destroy)

        def _export_image(self):
            p = filedialog.asksaveasfilename(defaultextension=".png",
                filetypes=[("PNG","*.png"),("JPEG","*.jpg")], title="Export Image", parent=self.win)
            if p: self.img.convert("RGB").save(p); self.status.set(f"Exported: {os.path.basename(p)}")

    # ===== PROCESS DOCUMENTER (main app window) =====

    class ProcessDocumenter:
        def __init__(self, root):
            self.root = root
            root.title("Auto-Doc")
            root.geometry("1060x960")
            root.configure(bg=T["BG"])
            root.resizable(True, True)
            self.steps = []
            self.step_counter = 1
            self.temp_dir = tempfile.mkdtemp()
            self.capture_paused = False
            self._pending_shot = None
            self.highlight_cursor = tk.BooleanVar(value=True)
            self.click_cap_var = tk.BooleanVar(value=False)
            self._mouse_listener = None
            self._last_focus_time = 0.0
            atexit.register(lambda: shutil.rmtree(self.temp_dir, ignore_errors=True))
            self._build_ui()
            root.bind("<FocusIn>", lambda e: setattr(self, "_last_focus_time", time.monotonic()))
            root.bind("<Map>", lambda e: setattr(self, "_last_focus_time", time.monotonic()))
            root.after(600, self._os_tips)

        def _build_ui(self):
            os_name = OS_NAMES.get(OS, OS)
            ob = {"Windows":"#1d4ed8","Darwin":"#7c3aed","Linux":"#15803d"}
            hdr = tk.Frame(self.root, bg=T["BG2"], pady=20); hdr.pack(fill="x")
            tk.Label(hdr, text="AUTO-DOC", font=(T["FONT"],20,"bold"), fg=T["ACCENT"], bg=T["BG2"]).pack()
            tk.Label(hdr, text="Process Documentation Tool", font=(T["FONT"],9),
                     fg=T["GREY"], bg=T["BG2"]).pack(pady=(2,6))
            br = tk.Frame(hdr, bg=T["BG2"]); br.pack()
            tk.Label(br, text=f"  {os_name}  ", font=(T["FONT"],8,"bold"), fg="white",
                     bg=ob.get(OS,"#374151"), padx=8, pady=3).pack(side="left")
            tk.Label(br, text="  Multi-Monitor  ·  Screenshot Editor  ·  Word Export",
                     font=(T["FONT"],9), fg=T["GREY"], bg=T["BG2"]).pack(side="left", padx=8)

            self.status_var = tk.StringVar(value=f"Ready  ·  {os_name}  ·  0 steps")
            tk.Label(self.root, textvariable=self.status_var, font=(T["FONT"],9), fg=T["ACCENT"],
                     bg=T["BG2"], anchor="w", padx=14, pady=7).pack(fill="x", side="bottom")

            style = ttk.Style(); style.theme_use("default")
            style.configure("D.TNotebook", background=T["BG"], borderwidth=0, tabmargins=[14,8,0,0])
            style.configure("D.TNotebook.Tab", background=T["BG2"], foreground=T["GREY"],
                            font=(T["FONT"],10,"bold"), padding=[18,8], borderwidth=0)
            style.map("D.TNotebook.Tab", background=[("selected",T["BG3"])], foreground=[("selected",T["ACCENT"])])

            nb = ttk.Notebook(self.root, style="D.TNotebook"); nb.pack(fill="both", expand=True)
            tab_main = tk.Frame(nb, bg=T["BG"]); nb.add(tab_main, text="  Documenter  ")
            tab_help = tk.Frame(nb, bg=T["BG"]); nb.add(tab_help, text="  How To Use  ")
            self._build_main_tab(tab_main); self._build_help_tab(tab_help)

        def _build_main_tab(self, parent):
            # Split layout: left (steps) + right (preview)
            pane = tk.PanedWindow(parent, orient="horizontal", bg=T["BG"], sashwidth=6, sashrelief="flat", bd=0)
            pane.pack(fill="both", expand=True)
            left = tk.Frame(pane, bg=T["BG"]); right = tk.Frame(pane, bg=T["BG"])
            pane.add(left, width=580, minsize=400); pane.add(right, width=440, minsize=250)

            # -- Left panel --
            info = self._section(left, "PROCESS INFO")
            self._field_row(info, "Process Title:", "title_var", width=38)
            self._field_row(info, "Author:        ", "author_var", width=38)

            entry = self._section(left, "ADD STEP")
            tk.Label(entry, text="Step description:", font=(T["FONT"],9), fg=T["GREY"], bg=T["BG3"]).pack(anchor="w")
            self.step_text = tk.Text(entry, height=3, font=(T["FONT"],10), bg=T["BG2"], fg=T["FG"],
                                     insertbackground=T["ACCENT"], relief="flat", bd=6, wrap="word")
            self.step_text.pack(fill="x", pady=(4,8))
            self.preview_var = tk.StringVar(value="")
            tk.Label(entry, textvariable=self.preview_var, font=(T["FONT"],8), fg="#ffeb3b", bg=T["BG3"]).pack(anchor="w")

            br = tk.Frame(entry, bg=T["BG3"]); br.pack(fill="x", pady=(4,0))
            make_button(br, "+ Add Step", self._add_step_no_shot,
                        "#14532d","#15803d", width=16).pack(side="left", padx=(0,8))
            make_button(br, "Screenshot + Add Step  (F9)", self._add_step_with_shot,
                        "#1d3461","#1d4ed8", width=28).pack(side="left")
            tk.Label(entry, text="  F9 = screenshot + add  ·  F10 = pause  ·  F11 = export",
                     font=(T["FONT"],8), fg=T["GREY"], bg=T["BG3"]).pack(anchor="w", pady=(6,0))
            orow = tk.Frame(entry, bg=T["BG3"]); orow.pack(anchor="w", pady=(4,0))
            _cb_kw = dict(font=(T["FONT"],9), fg="#80cbc4", bg=T["BG3"], selectcolor="#111c2b",
                          activebackground=T["BG3"], activeforeground="#80cbc4",
                          relief="flat", bd=0)
            tk.Checkbutton(orow, text="  Highlight cursor in screenshots",
                           variable=self.highlight_cursor, **_cb_kw).pack(side="left")
            tk.Checkbutton(orow, text="  Auto-capture on click",
                           variable=self.click_cap_var,
                           command=self._toggle_click_capture, **_cb_kw).pack(side="left", padx=(16,0))

            steps_sec = self._section(left, "DOCUMENTED STEPS  (double-click to edit screenshot)")
            ssb = tk.Scrollbar(steps_sec); ssb.pack(side="right", fill="y")
            self.steps_box = tk.Listbox(steps_sec, font=(T["FONT"],10), bg=T["BG2"], fg=T["FG"],
                                        selectbackground="#1d4ed8", selectforeground=T["FG"],
                                        relief="flat", bd=0, activestyle="none",
                                        yscrollcommand=ssb.set)
            self.steps_box.pack(fill="both", expand=True); ssb.config(command=self.steps_box.yview)
            self.steps_box.bind("<<ListboxSelect>>", self._on_step_select)
            self.steps_box.bind("<Double-1>", lambda e: self._open_editor())

            ar = tk.Frame(steps_sec, bg=T["BG3"]); ar.pack(fill="x", pady=(5,0))
            for lbl, cmd, bg, hbg in [
                ("Edit",     self._edit_step,    "#1e3a4a","#1e4d63"),
                ("Delete",   self._delete_step,  "#450a0a","#7f1d1d"),
                ("↑ Up",     self._move_up,      "#1e293b","#334155"),
                ("↓ Down",   self._move_down,    "#1e293b","#334155"),
                ("View/Edit",self._open_editor,  "#1d3461","#1d4ed8"),
                ("Save",     self._save_session, "#14532d","#15803d"),
                ("Load",     self._load_session, "#1e293b","#334155"),
                ("Export",   self._export_word,  "#1d4ed8","#2563eb"),
            ]:
                make_button(ar, lbl, cmd, bg, hbg).pack(side="left", padx=(0,3))

            # -- Right panel: preview --
            self._build_preview(right)

        def _build_preview(self, parent):
            make_section_label(parent, "SCREENSHOT PREVIEW")
            pc = tk.Frame(parent, bg=T["BG3"], padx=8, pady=8)
            pc.pack(fill="both", expand=True, padx=14, pady=(0,6))
            self.pinfo = tk.StringVar(value="Select a step to preview its screenshot")
            tk.Label(pc, textvariable=self.pinfo, font=(T["FONT"],8), fg=T["GREY"], bg=T["BG3"]).pack(anchor="w", pady=(0,6))
            self.pcanvas = tk.Canvas(pc, bg=T["BG2"], highlightthickness=0, bd=0)
            self.pcanvas.pack(fill="both", expand=True)
            self.pcanvas.bind("<Button-1>", lambda e: self._open_editor())
            self._ptk = None
            df = tk.Frame(pc, bg=T["BG3"]); df.pack(fill="x", pady=(6,0))
            tk.Label(df, text="Description:", font=(T["FONT"],8,"bold"), fg=T["GREY"], bg=T["BG3"]).pack(anchor="w")
            self.pdesc = tk.StringVar(value="")
            tk.Label(df, textvariable=self.pdesc, font=(T["FONT"],9), fg=T["FG"], bg=T["BG3"],
                     wraplength=380, justify="left", anchor="w").pack(fill="x")
            tk.Label(pc, text="Click preview or View/Edit to open editor",
                     font=(T["FONT"],8), fg="#37474f", bg=T["BG3"]).pack(anchor="w", pady=(6,0))

        def _on_step_select(self, e=None):
            sel = self.steps_box.curselection()
            if not sel: return
            step = self.steps[sel[0]]; self.pdesc.set(step.get("description",""))
            sp = step.get("screenshot_path","")
            if sp and os.path.exists(sp):
                try:
                    img = Image.open(sp); self.pcanvas.update_idletasks()
                    cw = max(self.pcanvas.winfo_width(), 200); ch = max(self.pcanvas.winfo_height(), 200)
                    scale = min(cw/img.width, ch/img.height, 1.0)
                    thumb = img.resize((max(1,int(img.width*scale)), max(1,int(img.height*scale))), Image.LANCZOS)
                    self._ptk = ImageTk.PhotoImage(thumb); self.pcanvas.delete("all")
                    self.pcanvas.create_image(cw//2, ch//2, anchor="center", image=self._ptk)
                    self.pinfo.set(f"Step {step['step_num']}  |  {img.width}x{img.height}  (click to edit)")
                except Exception as ex: self.pinfo.set(f"Error: {ex}"); self.pcanvas.delete("all")
            else:
                self.pcanvas.delete("all"); self.pinfo.set(f"Step {step['step_num']}  |  no screenshot")

        def _open_editor(self):
            sel = self.steps_box.curselection()
            if not sel: messagebox.showinfo("Nothing Selected","Click a step first."); return
            step = self.steps[sel[0]]; sp = step.get("screenshot_path","")
            if not sp or not os.path.exists(sp):
                messagebox.showinfo("No Screenshot", f"Step {step['step_num']} has no screenshot.")
                return
            try: img = Image.open(sp)
            except Exception as ex: messagebox.showerror("Error", str(ex)); return
            def on_save(edited):
                edited.save(sp, "PNG"); self._set_status(f"Step {step['step_num']} screenshot updated")
                self._on_step_select()
            ScreenshotEditor(self.root, img, on_save=on_save)

        # -- Help tab --
        def _build_help_tab(self, parent):
            canvas = tk.Canvas(parent, bg=T["BG"], highlightthickness=0, bd=0)
            sb = tk.Scrollbar(parent, orient="vertical", command=canvas.yview)
            inner = tk.Frame(canvas, bg=T["BG"])
            inner.bind("<Configure>", lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
            canvas.create_window((0,0), window=inner, anchor="nw")
            canvas.configure(yscrollcommand=sb.set)
            canvas.pack(side="left", fill="both", expand=True); sb.pack(side="right", fill="y")
            def _scroll(e): canvas.yview_scroll(int(-1*(e.delta/120)), "units")
            canvas.bind_all("<MouseWheel>", _scroll)
            PAD = 24
            def h1(t):
                tk.Label(inner, text=t, font=(T["FONT"],13,"bold"), fg=T["ACCENT"], bg=T["BG"], anchor="w").pack(fill="x", padx=PAD, pady=(20,4))
                tk.Frame(inner, bg="#1a3a4a", height=1).pack(fill="x", padx=PAD, pady=(0,8))
            def h2(t):
                tk.Label(inner, text=f"  {t}", font=(T["FONT"],10,"bold"), fg="#80cbc4", bg=T["BG"], anchor="w").pack(fill="x", padx=PAD, pady=(12,2))
            def body(t):
                tk.Label(inner, text=t, font=(T["FONT"],9), fg="#90a4ae", bg=T["BG"], anchor="w", justify="left", wraplength=720).pack(fill="x", padx=PAD+8, pady=1)
            def srow(n, title, desc):
                row = tk.Frame(inner, bg="#0d1f35", padx=14, pady=10); row.pack(fill="x", padx=PAD, pady=3)
                tk.Label(row, text=f" {n} ", font=(T["FONT"],11,"bold"), fg="#0a0f1a", bg=T["ACCENT"], padx=6, pady=2).pack(side="left", anchor="n", padx=(0,12))
                col = tk.Frame(row, bg="#0d1f35"); col.pack(side="left", fill="x", expand=True)
                tk.Label(col, text=title, font=(T["FONT"],10,"bold"), fg=T["FG"], bg="#0d1f35", anchor="w").pack(fill="x")
                tk.Label(col, text=desc, font=(T["FONT"],9), fg="#607d8b", bg="#0d1f35", anchor="w", justify="left", wraplength=640).pack(fill="x")
            def div(): tk.Frame(inner, bg="#1a2a3a", height=1).pack(fill="x", padx=PAD, pady=8)

            h1("  HOW TO USE THIS TOOL")
            srow("1","Fill In Process Info","Enter a Process Title and Author for the exported document.")
            srow("2","Type Your Step Description","Describe the action you performed or are about to perform.")
            srow("3a","Screenshot + Add Step (F9) - recommended",
                 "Type description, then press F9. The app minimises, captures the monitor your mouse is on, and saves both together.")
            srow("3b","Add Step Without Screenshot","Click + Add Step for text-only steps.")
            srow("4","Review & Annotate Screenshots",
                 "Select a step to see its thumbnail on the right. Double-click or press View/Edit to open the full editor. "
                 "Annotate with arrows, rectangles, freehand drawing, text, or highlighter. Crop unwanted areas.")
            srow("5","Edit, Reorder, Save","Use Edit, Delete, Up, Down to manage steps. Save/Load for session persistence.")
            srow("6","Export to Word (F11)","A formatted .docx is saved to ~/Documents and opened automatically.")
            div()
            h1("  SCREENSHOT EDITOR TOOLS")
            body("Arrow - click and drag to draw an arrow pointing at something")
            body("Rectangle - draw a box around an area of interest")
            body("Freehand - draw freely with the mouse like a pen")
            body("Text - click to place a text annotation at that spot")
            body("Highlight - semi-transparent yellow box like a marker")
            body("Crop - drag a region to crop the image to that area")
            body("Also: Color picker, Width slider (1-12px), Undo (30 levels), Reset, Save to Step, Export Image")
            div()
            h1("  KEYBOARD SHORTCUTS")
            body("F9 = Screenshot + Add Step  |  F10 = Pause/Resume  |  F11 = Export to Word")
            body("These work globally, even when the app is minimised.")
            div()
            h1("  MULTI-MONITOR SUPPORT")
            body("Only the monitor your mouse cursor is on gets captured. Move your mouse to the target screen before pressing F9.")
            if OS in ("Darwin","Linux"):
                div(); h1("  PLATFORM NOTES")
                if OS == "Darwin": body("Grant Screen Recording in System Settings > Privacy & Security > Screen Recording.")
                else: body("Install xdotool: sudo apt install xdotool. If F-keys fail: sudo usermod -aG input $USER")
            tk.Frame(inner, bg=T["BG"], height=30).pack()

        # -- Helpers --
        def _section(self, parent, title):
            outer = tk.Frame(parent, bg=T["BG"])
            is_steps = title.startswith("DOCUMENTED STEPS")
            outer.pack(fill="both" if is_steps else "x", expand=is_steps)
            lf = tk.Frame(outer, bg=T["BG"]); lf.pack(fill="x", padx=14, pady=(10,2))
            tk.Label(lf, text=f"  {title}  ", font=(T["FONT"],8,"bold"), fg=T["ACCENT"], bg=T["BG"]).pack(side="left")
            tk.Frame(lf, bg=T["BORDER"], height=1).pack(side="left", fill="x", expand=True, pady=6)
            inner = tk.Frame(outer, bg=T["BG3"], padx=12, pady=10)
            inner.pack(fill="both", expand=True, padx=14, pady=(0,8))
            return inner

        def _field_row(self, parent, label, varname, width=40):
            row = tk.Frame(parent, bg=T["BG3"]); row.pack(fill="x", pady=4)
            tk.Label(row, text=label, font=(T["FONT"],9), fg=T["GREY"], bg=T["BG3"]).pack(side="left")
            var = tk.StringVar(); setattr(self, varname, var)
            tk.Entry(row, textvariable=var, font=(T["FONT"],10), bg=T["BG2"], fg=T["FG"],
                     insertbackground=T["ACCENT"], relief="flat", bd=6, width=width).pack(side="left", padx=6)

        def _set_status(self, msg):
            p = "  ·  PAUSED" if self.capture_paused else ""
            self.status_var.set(f"{msg}  ·  {len(self.steps)} steps{p}")

        def _os_tips(self):
            tips = {"Darwin":("macOS Permissions","For screenshots:\nSystem Settings > Privacy & Security > Screen Recording\n> Enable for Terminal. Restart app after granting."),
                    "Linux":("Linux Tips","Install xdotool: sudo apt install xdotool\nIf F-keys fail: sudo usermod -aG input $USER")}
            if OS in tips: messagebox.showinfo(*tips[OS])

        # -- Click capture --
        def _is_in_app_window(self, x, y):
            try:
                wx, wy = self.root.winfo_rootx(), self.root.winfo_rooty()
                ww, wh = self.root.winfo_width(), self.root.winfo_height()
                return wx <= x <= wx + ww and wy <= y <= wy + wh
            except:
                return False

        def _toggle_click_capture(self):
            from pynput.mouse import Listener as MouseListener, Button
            if self.click_cap_var.get():
                def on_click(x, y, button, pressed):
                    if not pressed or button != Button.left:
                        return
                    if self.capture_paused or not self.click_cap_var.get():
                        return
                    if self._is_in_app_window(x, y):
                        return
                    self.root.after(150, self._auto_capture_step)
                self._mouse_listener = MouseListener(on_click=on_click)
                self._mouse_listener.daemon = True
                self._mouse_listener.start()
                self._set_status("Auto-capture on click: ON")
            else:
                if self._mouse_listener:
                    try: self._mouse_listener.stop()
                    except: pass
                    self._mouse_listener = None
                self._set_status("Auto-capture on click: OFF")

        def _auto_capture_step(self):
            if self.capture_paused or not self.click_cap_var.get():
                return
            if time.monotonic() - self._last_focus_time < 1.0:
                return  # Ignore the click that restored the app
            # App is already minimized — take screenshot silently without deiconifying
            img = take_screenshot_at_mouse(highlight=self.highlight_cursor.get())
            path = None
            if img:
                try: path = self._save_screenshot(img)
                except: pass
            desc = self.step_text.get("1.0", "end").strip() or f"Step {self.step_counter}"
            self.step_text.delete("1.0", "end")
            self._commit_step(desc, path)

        # -- Screenshot actions --
        def _do_screenshot(self):
            if self.capture_paused: return None
            self.root.iconify(); self.root.update(); time.sleep(0.35)
            img = take_screenshot_at_mouse(highlight=self.highlight_cursor.get())
            self.root.deiconify(); self.root.lift(); return img

        def _save_screenshot(self, img):
            p = os.path.join(self.temp_dir, f"step_{self.step_counter}_{datetime.datetime.now().strftime('%H%M%S%f')}.png")
            img.save(p, "PNG"); return p

        def _toggle_pause_capture(self):
            self.capture_paused = not self.capture_paused
            if self.capture_paused:
                self.preview_var.set("  PAUSED  (F10 to resume)"); self._set_status("Screenshot capture paused")
            else: self.preview_var.set(""); self._set_status("Ready")

        # -- Step operations --
        def _get_description(self):
            desc = self.step_text.get("1.0","end").strip()
            if not desc: messagebox.showwarning("Empty","Please type a step description."); return None
            return desc

        def _add_step_no_shot(self):
            desc = self._get_description()
            if desc is None: return
            path = None
            if self._pending_shot:
                try: path = self._save_screenshot(self._pending_shot)
                except: pass
                self._pending_shot = None; self.preview_var.set("")
            self._commit_step(desc, path)

        def _add_step_with_shot(self):
            desc = self._get_description()
            if desc is None: return
            img = self._do_screenshot(); path = None
            if img:
                try: path = self._save_screenshot(img)
                except: pass
            self._pending_shot = None; self.preview_var.set("")
            self._commit_step(desc, path)

        def _commit_step(self, desc, spath):
            step = {"step_num": self.step_counter, "description": desc,
                    "screenshot_path": spath,
                    "timestamp": datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")}
            self.steps.append(step); self.step_counter += 1
            self.step_text.delete("1.0","end"); self._pending_shot = None; self.preview_var.set("")
            self._refresh_list()
            icon = "IMG" if spath else "TXT"
            self._set_status(f"{icon} Step {step['step_num']} added")

        def _refresh_list(self):
            self.steps_box.delete(0,"end")
            for s in self.steps:
                icon = "IMG" if s.get("screenshot_path") else "TXT"
                self.steps_box.insert("end", f"  [{icon}]  {s['step_num']:>2}.  {s['description'][:70]}")

        def _selected_idx(self):
            sel = self.steps_box.curselection()
            if not sel: messagebox.showinfo("Nothing Selected","Click a step first."); return None
            return sel[0]

        def _edit_step(self):
            idx = self._selected_idx()
            if idx is None: return
            step = self.steps[idx]
            new = simpledialog.askstring("Edit Step", f"Edit Step {step['step_num']}:",
                                         initialvalue=step["description"], parent=self.root)
            if new and new.strip(): step["description"] = new.strip(); self._refresh_list()

        def _delete_step(self):
            idx = self._selected_idx()
            if idx is None: return
            if messagebox.askyesno("Delete", f"Delete Step {self.steps[idx]['step_num']}?"):
                self.steps.pop(idx); self._renumber(); self._refresh_list(); self._set_status("Step deleted")

        def _move_up(self):
            idx = self._selected_idx()
            if idx is None or idx == 0: return
            self.steps[idx-1], self.steps[idx] = self.steps[idx], self.steps[idx-1]
            self._renumber(); self._refresh_list(); self.steps_box.selection_set(idx-1)

        def _move_down(self):
            idx = self._selected_idx()
            if idx is None or idx >= len(self.steps)-1: return
            self.steps[idx+1], self.steps[idx] = self.steps[idx], self.steps[idx+1]
            self._renumber(); self._refresh_list(); self.steps_box.selection_set(idx+1)

        def _renumber(self):
            for i, s in enumerate(self.steps): s["step_num"] = i+1
            self.step_counter = len(self.steps)+1

        # -- Session save/load --
        def _save_session(self):
            path = filedialog.asksaveasfilename(defaultextension=".json",
                filetypes=[("Word Document","*.docx"),("JSON","*.json")], title="Save / Export")
            if not path: return
            if path.lower().endswith(".docx"):
                self._export_word(out_path=path); return
            data = {"title": self.title_var.get(), "author": self.author_var.get(),
                    "steps": [{k:v for k,v in s.items() if k != "screenshot_path"} for s in self.steps]}
            with open(path,"w") as f: json.dump(data, f, indent=2)
            messagebox.showinfo("Saved", f"Session saved:\n{path}")

        def _load_session(self):
            path = filedialog.askopenfilename(filetypes=[("JSON","*.json")])
            if not path: return
            with open(path) as f: data = json.load(f)
            self.title_var.set(data.get("title","")); self.author_var.set(data.get("author",""))
            self.steps = data.get("steps",[]); self._renumber(); self._refresh_list()
            self._set_status("Session loaded")

        # -- Export to Word --
        def _export_word(self, out_path=None):
            if not self.steps: messagebox.showwarning("No Steps","No steps to export."); return
            title = self.title_var.get().strip() or "Auto-Doc"
            author = self.author_var.get().strip() or "IT Service Desk"
            date_s = datetime.datetime.now().strftime("%B %d, %Y  %H:%M")
            os_s = OS_NAMES.get(OS, OS)

            doc = Document()
            for sec in doc.sections:
                sec.top_margin = sec.bottom_margin = Inches(1)
                sec.left_margin = sec.right_margin = Inches(1.2)

            def srun(para, text, size=10, bold=False, italic=False, color=None):
                r = para.add_run(text); r.font.name = "Calibri"; r.font.size = Pt(size)
                r.font.bold = bold; r.font.italic = italic
                if color: r.font.color.rgb = RGBColor(*color)
                return r

            def rule(color="1A2A3A"):
                p = doc.add_paragraph(); pPr = p._p.get_or_add_pPr()
                pBdr = OxmlElement("w:pBdr"); bot = OxmlElement("w:bottom")
                bot.set(qn("w:val"),"single"); bot.set(qn("w:sz"),"6")
                bot.set(qn("w:space"),"1"); bot.set(qn("w:color"),color)
                pBdr.append(bot); pPr.append(pBdr)

            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srun(p, "AUTO-DOC  |  PROCESS DOCUMENTATION", size=9, bold=True, color=(0x00,0x96,0x88))
            p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srun(p2, title, size=24, bold=True, color=(0x01,0x27,0x9A))
            p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srun(p3, f"Author: {author}  |  {date_s}  |  {len(self.steps)} Steps  |  {os_s}",
                 size=10, color=(0x60,0x7D,0x8B))
            rule("0D47A1"); doc.add_paragraph()

            for step in self.steps:
                sp = doc.add_paragraph()
                srun(sp, f"Step {step['step_num']}  --  ", size=13, bold=True, color=(0x01,0x27,0x9A))
                srun(sp, step["description"], size=13, color=(0x21,0x21,0x21))
                tp = doc.add_paragraph()
                srun(tp, f"  Recorded: {step['timestamp']}", size=9, italic=True, color=(0x90,0xA4,0xAE))

                spath = step.get("screenshot_path","")
                if spath and os.path.exists(spath):
                    try:
                        img = Image.open(spath); w,h = img.size; max_w = 6.0; w_in = w/96
                        if w_in > max_w:
                            sc = max_w/w_in; img = img.resize((int(w*sc),int(h*sc)), Image.LANCZOS)
                            buf = io.BytesIO(); img.save(buf,"PNG"); buf.seek(0)
                            doc.add_picture(buf, width=Inches(max_w))
                        else: doc.add_picture(spath, width=Inches(w_in))
                        cp = doc.add_paragraph(); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        srun(cp, f"Screenshot -- Step {step['step_num']}", size=9, italic=True, color=(0x78,0x90,0x9C))
                    except: pass
                doc.add_paragraph()

            rule("263545")
            fp = doc.add_paragraph(); fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            srun(fp, f"Generated by Auto-Doc  |  {os_s}  |  {date_s}", size=9, color=(0x78,0x90,0x9C))

            safe = "".join(c for c in title if c.isalnum() or c in " _-").strip().replace(" ","_")
            fname = f"{safe}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            if not out_path:
                out_path = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word Document", "*.docx")],
                    initialdir=os.path.expanduser("~/Documents"),
                    initialfile=fname,
                    title="Export to Word",
                )
            if not out_path: return
            doc.save(out_path); messagebox.showinfo("Exported!", f"Saved to:\n{out_path}")
            open_file(out_path)

    # -- Run --
    root = tk.Tk(); app = ProcessDocumenter(root); _app_ref[0] = app
    def on_close():
        if fkey_listener:
            try: fkey_listener.stop()
            except: pass
        if app._mouse_listener:
            try: app._mouse_listener.stop()
            except: pass
        root.destroy()
    root.protocol("WM_DELETE_WINDOW", on_close); root.mainloop()

# === ENTRY POINT ===
if __name__ == "__main__":
    if PY_VERSION < (3, 8):
        print(f"\nPython 3.8+ required. You have {PY_VERSION.major}.{PY_VERSION.minor}.")
        sys.exit(1)
    if "--run-app" in sys.argv: _run_main_app()
    else: root = tk.Tk(); Launcher(root); root.mainloop()
